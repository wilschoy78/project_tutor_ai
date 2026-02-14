import re
import io
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None
    print("Warning: pypdf not installed. PDF parsing will be disabled.")

try:
    import docx
except ImportError:
    docx = None
    print("Warning: python-docx not installed. DOCX parsing will be disabled.")

from typing import List, Dict, Any
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings, ChatOllama
from langchain_community.embeddings import FastEmbedEmbeddings # Lightweight CPU embeddings
try:
    from langchain_mistralai import MistralAIEmbeddings, ChatMistralAI
except ImportError:
    print("Warning: langchain-mistralai not installed. Mistral API mode will fail if selected.")

try:
    from langchain_groq import ChatGroq
    # Groq doesn't have native embeddings yet, usually uses HuggingFace or others.
    # We will use OllamaEmbeddings or fallback to a lightweight one if needed.
except ImportError:
    print("Warning: langchain-groq not installed. Groq mode will fail if selected.")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.schema import Document
from langchain.prompts import PromptTemplate
from app.core.config import settings
from app.services.moodle_client import moodle_client
from app.services.student_service import student_service

class RAGService:
    def __init__(self):
        # Initialize LLM and Embeddings based on Provider
        if settings.LLM_PROVIDER == "mistral_api":
            if not settings.MISTRAL_API_KEY:
                raise ValueError("MISTRAL_API_KEY is required when using 'mistral_api' provider.")
            
            print("Initializing RAG with Mistral AI API...")
            self.embeddings = MistralAIEmbeddings(
                mistral_api_key=settings.MISTRAL_API_KEY,
                model="mistral-embed"
            )
            self.llm = ChatMistralAI(
                mistral_api_key=settings.MISTRAL_API_KEY,
                model=settings.MODEL_NAME, # e.g., "mistral-small-latest"
                temperature=0.7
            )
        elif settings.LLM_PROVIDER == "groq":
            if not settings.GROQ_API_KEY:
                raise ValueError("GROQ_API_KEY is required when using 'groq' provider.")
            
            print("Initializing RAG with Groq & FastEmbed (Lightweight CPU) Embeddings...")
            # Groq is purely for Inference (LLM), not Embeddings.
            # We use FastEmbedEmbeddings which is extremely lightweight and faster than HuggingFace/PyTorch.
            # This allows deployment on 512MB RAM instances (like Render Free Tier).
            
            self.embeddings = FastEmbedEmbeddings(
                model_name="BAAI/bge-small-en-v1.5"
            )
            
            self.llm = ChatGroq(
                groq_api_key=settings.GROQ_API_KEY,
                model_name=settings.MODEL_NAME, # e.g., "llama3-8b-8192"
                temperature=0.7
            )
        else:
            # Default to Ollama (Local)
            print("Initializing RAG with Local Ollama...")
            self.embeddings = OllamaEmbeddings(
                base_url=settings.OLLAMA_BASE_URL,
                model=settings.MODEL_NAME
            )
            self.llm = ChatOllama(
                base_url=settings.OLLAMA_BASE_URL,
                model=settings.MODEL_NAME,
                temperature=0.7
            )
        
        # Initialize Vector Store (ChromaDB)
        # Persistent storage in ./chroma_db
        self.vector_store = Chroma(
            persist_directory="./chroma_db",
            embedding_function=self.embeddings,
            collection_name="moodle_content"
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )

    def _clean_html(self, raw_html: str) -> str:
        """Helper to strip HTML tags from Moodle content."""
        if not raw_html:
            return ""
        cleanr = re.compile('<.*?>')
        cleantext = re.sub(cleanr, '', raw_html)
        return cleantext.strip()

    def ingest_course_content(self, course_id: int) -> Dict[str, Any]:
        """
        Fetches content from Moodle (or Mock), chunks it, and stores in Vector DB.
        """
        print(f"Ingesting content for course {course_id}...")
        
        # 0. Clear existing content for this course to prevent duplicates
        try:
            # Note: Chroma's delete might throw if collection is empty or other issues, so we wrap in try/except
            print(f"Clearing existing vector data for course {course_id}...")
            # We need to get ids to delete, or use where clause if supported by the wrapper
            # LangChain's Chroma wrapper .delete() expects IDs. 
            # So we first get all IDs for this course.
            
            # Using the underlying client to get IDs
            existing_docs = self.vector_store.get(where={"course_id": course_id})
            if existing_docs and existing_docs['ids']:
                print(f"Deleting {len(existing_docs['ids'])} existing documents...")
                self.vector_store.delete(ids=existing_docs['ids'])
                print("Deletion complete.")
            else:
                print("No existing documents found for this course.")
                
        except Exception as e:
            print(f"Warning during cleanup: {e}")

        # 1. Fetch content
        contents = moodle_client.get_course_contents(course_id)
        
        # 1.5 Fetch User Activities (Grades & Completion) - NEW FEATURE
        # Note: In a real multi-user RAG, you might not want to ingest specific student grades into the GLOBAL vector store 
        # because that would make Student A's grades visible to Student B via RAG.
        # HOWEVER, if this is for the *Teacher's* Knowledge Base (to ask "Who failed the quiz?"), then it makes sense.
        # We will ingest it but label it clearly. 
        # CAUTION: For privacy, ensure the RAG prompt respects user roles, or only ingest this for the teacher's view.
        # For now, we will fetch generic activity structure, not individual student grades for the RAG.
        # If the user wants to ingest *aggregated* stats, that's safer.
        
        documents = []
        
        # 2. Process content into Documents
        for section in contents:
            section_name = section.get("name", "Unnamed Section")
            for module in section.get("modules", []):
                mod_name = module.get("name", "Unnamed Module")
                mod_type = module.get("modname", "unknown")
                
                # Start building text content
                content_text = f"Course ID: {course_id}\nSection: {section_name}\nModule: {mod_name}\nType: {mod_type}\n"
                
                # 1. Get Description (common for all modules)
                if "description" in module and module["description"]:
                    desc_clean = self._clean_html(module["description"])
                    if desc_clean:
                        content_text += f"Description: {desc_clean}\n"
                
                # Check for Quiz-specific data if available in course contents (usually limited)
                if mod_type == "quiz":
                    # In standard core_course_get_contents, quiz details are minimal (just intro/dates).
                    # Deep quiz question extraction requires mod_quiz_get_quizzes_by_courses, 
                    # but that's a separate API call we might add later if needed.
                    # For now, we rely on the description and any attached files.
                    if "dates" in module:
                        content_text += f"Dates: {module['dates']}\n"
                
                # 2. Get Page Content (specific to 'page' modname) or generic 'contents'
                # Moodle returns a list 'contents' for resources/pages
                if "contents" in module:
                    for item in module["contents"]:
                        # 'content' field usually holds the HTML for Pages
                        if "content" in item:
                            page_clean = self._clean_html(item["content"])
                            if page_clean:
                                content_text += f"Content: {page_clean}\n"
                        
                        # 'filename' is useful for Resources (PDFs, etc.)
                        if "filename" in item:
                             content_text += f"File Attachment: {item['filename']}\n"
                             
                             # PDF Extraction Logic
                             is_pdf = False
                             if "mimetype" in item and item["mimetype"] == "application/pdf":
                                 is_pdf = True
                             elif item['filename'].lower().endswith(".pdf"):
                                 is_pdf = True
                                 
                             if is_pdf and "fileurl" in item:
                                 if PdfReader:
                                     try:
                                         print(f"Downloading PDF: {item['filename']}")
                                         file_bytes = moodle_client.download_file(item["fileurl"])
                                         if file_bytes:
                                             file_obj = io.BytesIO(file_bytes)
                                             reader = PdfReader(file_obj)
                                             pdf_text = ""
                                             for page in reader.pages:
                                                 extracted = page.extract_text()
                                                 if extracted:
                                                     pdf_text += extracted + "\n"
                                             
                                             if pdf_text:
                                                 content_text += f"--- PDF CONTENT START ({item['filename']}) ---\n{pdf_text}\n--- PDF CONTENT END ---\n"
                                                 print(f"Extracted {len(pdf_text)} chars from PDF.")
                                             else:
                                                 print("PDF was empty or unreadable.")
                                     except Exception as e:
                                         print(f"Error parsing PDF {item['filename']}: {e}")
                                         content_text += f"[Error reading PDF content: {str(e)}]\n"
                                 else:
                                     content_text += "[PDF content not extracted: pypdf library missing]\n"
                             
                             # DOCX Extraction Logic
                             is_docx = False
                             if "mimetype" in item and item["mimetype"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                 is_docx = True
                             elif item['filename'].lower().endswith(".docx"):
                                 is_docx = True
                                 
                             if is_docx and "fileurl" in item:
                                 if docx:
                                     try:
                                         print(f"Downloading DOCX: {item['filename']}")
                                         file_bytes = moodle_client.download_file(item["fileurl"])
                                         if file_bytes:
                                             file_obj = io.BytesIO(file_bytes)
                                             doc_file = docx.Document(file_obj)
                                             docx_text = "\n".join([para.text for para in doc_file.paragraphs])
                                             
                                             if docx_text:
                                                 content_text += f"--- DOCX CONTENT START ({item['filename']}) ---\n{docx_text}\n--- DOCX CONTENT END ---\n"
                                                 print(f"Extracted {len(docx_text)} chars from DOCX.")
                                             else:
                                                 print("DOCX was empty or unreadable.")
                                     except Exception as e:
                                         print(f"Error parsing DOCX {item['filename']}: {e}")
                                         content_text += f"[Error reading DOCX content: {str(e)}]\n"
                                 else:
                                     content_text += "[DOCX content not extracted: python-docx library missing]\n"
                             
                             # DOC Extraction Logic (Warning)
                             if item['filename'].lower().endswith(".doc"):
                                 content_text += f"[WARNING: .doc file ({item['filename']}) skipped. Please convert to .docx or PDF for AI ingestion.]\n"

                doc = Document(
                    page_content=content_text,
                    metadata={
                        "course_id": course_id,
                        "source": f"{section_name} - {mod_name}",
                        "type": mod_type,
                        "module": mod_name,
                        "section": section_name
                    }
                )
                documents.append(doc)
        
        if not documents:
            print("No documents found to ingest.")
            return {"status": "warning", "message": "No content found"}

        # 3. Split and Store
        chunks = self.text_splitter.split_documents(documents)
        self.vector_store.add_documents(chunks)
        self.vector_store.persist()
        
        print(f"Ingested {len(chunks)} chunks for course {course_id}")
        return {"status": "success", "chunks_count": len(chunks)}

    def ask_question(self, course_id: int, question: str, student_id: int = 1):
        """
        RAG Pipeline: Retrieve relevant docs -> Generate Answer
        """
        # 1. Get Student Context
        profile = student_service.get_student_profile(student_id)
        progress = student_service.get_student_progress(student_id, course_id)
        
        # Format quiz scores safely to avoid curly braces in PromptTemplate
        quiz_scores_str = ", ".join([f"{k}: {v}" for k, v in progress['quiz_scores'].items()]) if progress['quiz_scores'] else "None"
        
        student_context = f"""
        Student Profile:
        - Name: {profile['name']}
        - Learning Style: {profile['learning_style']}
        - Strengths: {', '.join(profile['strengths'])}
        - Weaknesses: {', '.join(profile['weaknesses'])}
        
        Student Progress:
        - Completed: {', '.join(progress['completed_modules'])}
        - Quiz Scores: {quiz_scores_str}
        """
        
        # 2. Define Custom Prompt
        template = """
        You are an AI Tutor personalized for a specific student. 
        Use the following pieces of context (Course Material) to answer the question at the end.
        
        Course Material:
        {context}
        
        Target Student Context:
        """ + student_context + """
        
        Instructions:
        - Adapt your explanation to the student's learning style.
        - If the student is weak in a topic, provide extra examples.
        - Reference their progress if relevant (e.g., "Recall from Week 1...").
        - If you don't know the answer, just say that you don't know, don't try to make up an answer.
        
        Question: {question}
        Helpful Answer:
        """
        
        QA_CHAIN_PROMPT = PromptTemplate.from_template(template)

        # 3. Create Retriever
        retriever = self.vector_store.as_retriever(
            search_kwargs={
                "k": 3,
                "filter": {"course_id": course_id}
            }
        )
        
        # 4. Create QA Chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=retriever,
            return_source_documents=True,
            chain_type_kwargs={"prompt": QA_CHAIN_PROMPT}
        )
        
        # 5. Execute
        result = qa_chain.invoke({"query": question})
        
        return {
            "answer": result["result"],
            "sources": [doc.metadata for doc in result["source_documents"]]
        }

    def generate_quiz(self, course_id: int, topic: str):
        """
        Generates a multiple choice question based on the topic and course content.
        """
        # 1. Retrieve content
        retriever = self.vector_store.as_retriever(
            search_kwargs={
                "k": 3,
                "filter": {"course_id": course_id}
            }
        )
        docs = retriever.invoke(topic)
        context_text = "\n\n".join([doc.page_content for doc in docs])
        
        # 2. Prompt for Quiz Generation
        template = """
        You are an AI Tutor. Based on the following course content, generate a multiple-choice question to test the student's understanding of "{topic}".
        
        Course Content:
        {context}
        
        Output Format:
        You must return a valid JSON object with the following structure:
        {{
            "question": "The question text",
            "options": ["Option A", "Option B", "Option C", "Option D"],
            "correct_answer": "The correct option text (must be one of the options)",
            "explanation": "Brief explanation of why it is correct"
        }}
        
        Ensure the JSON is valid and has no markdown formatting (like ```json). Just the raw JSON string.
        """
        
        prompt = PromptTemplate(
            template=template,
            input_variables=["topic", "context"]
        )
        
        chain = prompt | self.llm
        
        # 3. Execute
        response = chain.invoke({"topic": topic, "context": context_text})
        
        # 4. Parse JSON
        import json
        try:
            content = response.content.strip()
            # Remove markdown code blocks if present
            if content.startswith("```json"):
                content = content[7:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()
            
            quiz_data = json.loads(content)
            return quiz_data
        except json.JSONDecodeError:
            # Fallback if JSON parsing fails
            print(f"Failed to parse quiz JSON: {response.content}")
            return {
                "question": f"Could not generate a structured quiz for {topic}. Please try again.",
                "options": ["Error"],
                "correct_answer": "Error",
                "explanation": "The model failed to produce valid JSON."
            }

    def generate_study_plan(self, course_id: int, weaknesses: List[str]):
        """
        Generates a personalized study plan based on identified weaknesses.
        """
        if not weaknesses:
            return "Great job! You seem to be doing well in all topics. Keep reviewing the latest materials."
            
        # 1. Retrieve relevant materials for weaknesses
        context_docs = []
        for topic in weaknesses:
            docs = self.vector_store.similarity_search(
                topic, 
                k=2,
                filter={"course_id": course_id}
            )
            context_docs.extend(docs)
            
        # Deduplicate docs based on content
        seen = set()
        unique_docs = []
        for doc in context_docs:
            if doc.page_content not in seen:
                seen.add(doc.page_content)
                unique_docs.append(doc)
        
        context_text = "\n\n".join([d.page_content for d in unique_docs])
        
        # 2. Prompt LLM
        prompt = f"""
        You are an expert Educational Planner. 
        The student has shown weaknesses in the following topics: {', '.join(weaknesses)}.
        
        Using the available course materials below, create a structured 3-step study plan to help them improve.
        For each step, recommend specific modules or concepts to review.
        
        Course Material:
        {context_text}
        
        Study Plan:
        """
        
        # Direct LLM call
        response = self.llm.invoke(prompt)
        
        if hasattr(response, 'content'):
            return response.content
        return str(response)

    def generate_learning_path(self, course_id: int, student_id: int):
        """
        Analyzes student performance and generates a personalized study path.
        """
        # 1. Get Progress
        progress = student_service.get_student_progress(student_id, course_id)
        quiz_scores = progress.get('quiz_scores', {})
        
        # 2. Identify Weaknesses (Score < 75)
        weaknesses = []
        for quiz, score in quiz_scores.items():
            if score < 75:
                # Extract topic from quiz name
                topic = quiz.replace(" Quiz", "").replace("Test", "").strip()
                weaknesses.append(topic)
                
        if not weaknesses:
            if not quiz_scores:
                # No data yet
                return {
                    "status": "start",
                    "message": "Welcome! Start by exploring the Course Introduction.",
                    "recommendations": ["Review Course Introduction"]
                }
            else:
                return {
                    "status": "on_track",
                    "message": "Excellent work! You are performing well in all assessed topics.",
                    "recommendations": ["Continue to the next module."]
                }

        # 3. Generate Plan
        plan_text = self.generate_study_plan(course_id, weaknesses)
        
        return {
            "status": "needs_improvement",
            "weaknesses": weaknesses,
            "study_plan": plan_text
        }

    def get_knowledge_base_summary(self, course_id: int):
        """
        Returns a summary of all ingested documents for a course.
        """
        try:
            # Query all documents for this course
            # Note: We use the underlying collection directly for metadata access if possible,
            # or rely on the wrapper. The LangChain wrapper doesn't expose metadata query easily
            # without fetching content.
            
            # Using the underlying Chroma client if available, or just the wrapper
            # The wrapper self.vector_store is a Chroma object.
            
            # Use get() method of the underlying collection
            result = self.vector_store.get(where={"course_id": course_id})
            
            if not result or not result['ids']:
                return {"course_id": course_id, "document_count": 0, "sources": []}
            
            # Process metadata to group by source
            sources = {}
            for i, meta in enumerate(result['metadatas']):
                # Construct a unique key for the source (e.g., "Module Name (Type)")
                mod_name = meta.get('module', 'Unknown Module')
                mod_type = meta.get('type', 'unknown')
                key = f"{mod_name} ({mod_type})"
                
                if key not in sources:
                    sources[key] = {
                        "name": mod_name,
                        "type": mod_type,
                        "chunks": 0,
                        "section": meta.get('section', 'Unknown Section')
                    }
                sources[key]["chunks"] += 1
                
            return {
                "course_id": course_id,
                "document_count": len(result['ids']),
                "sources": list(sources.values())
            }
            
        except Exception as e:
            print(f"Error getting knowledge base: {e}")
            return {"error": str(e)}

rag_service = RAGService()
